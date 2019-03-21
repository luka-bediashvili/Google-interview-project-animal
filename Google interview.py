#Create a catalogue (like encyclopedia) of data on adoptable pets (name, description, location,
#etc.) where one can browse and search the dataset. You can start with this dataset, modify it, or
#use another dataset with similar fields.
#1. Example search queries you'll want to be able to handle:
#○ all dogs
#○ by pet's name
#○ female cats with tabby colouring
#2. Add search autocomplete
#3. Bonus points for making it responsive and work well on small screens (phone, tablets)   (look into)
#and with touch interactions.
#4. Extra bonus points if you add animal images to the catalogue.
from tkinter import *
import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
import smtplib
import sqlite3
from PIL import Image, ImageTk
from tkinter import messagebox
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import hashlib
class Login(tk.Frame):
    def __init__(self, parent):
        tk.Frame.__init__(self, parent)
        Lable1 = tk.Label(self, text="User ID")
        Lable1.grid(row=1, column=1, sticky=W)

        self.user_name = ttk.Entry(self)
        self.user_name.grid(row=1, column=3, sticky=W)

        Lable1 = tk.Label(self, text="Password")
        Lable1.grid(row=2, column=1, sticky=W)

        self.password = ttk.Entry(self,show="*")
        self.password.grid(row=2, column=3, sticky=W)

        button1 = ttk.Button(self, text="Enter", command=lambda: self.UserLogin())
        button1.grid(row=3, column=2, sticky=W)
        button2 = ttk.Button(self, text="Join the community", command=lambda: self.create_user_open())
        button2.grid(row=4, column=2, sticky=W)
    def UserLogin(self):
        #with sqlite3.connect('Data.db') as db:
        #    data = db.cursor()
        #    data.execute(
        #    'select * from "User Data" where Password=? and User_ID=?', (hashlib.md5(self.password.get().encode('utf-8')).hexdigest(),self.user_name.get()))
        #self.check = data.fetchall()
        #if len(self.check)>0:
        #    global admin
        #    admin = self.check[0][5]
        #    print(admin)
        open.destroy()
        window = Tk()
        window.title('Animal Place')
        Main_body(window).grid()
        window.mainloop()
        #else:
        #    tk.messagebox.showinfo("ERROR", "Incorrect username or password")
    def create_user_open(self):
        create_user_open=Tk()
        create_user_open.title('New User')
        create_user(create_user_open).grid()
        create_user_open.mainloop()
class Main_body(tk.Frame):
    def __init__(self, parent):
        tk.Frame.__init__(self, parent)
        self.frame=Frame(parent)
        self.frame.grid()
        self.text=''
        self.text_show=StringVar()
        self.text_show.set(self.text)
        self.bind("<KeyPress>",self.keydown)
        self.grid()
        self.focus_set()
        self.data=[]
        self.count = 5
        try:
            Lable1 = tk.Label(self.frame, text='Search for: ')
            Lable1.grid(row=1, column=1, sticky=W)

            #self.user_input = Label(self.frame, textvariable = self.text_show)
            #self.user_input.grid(row=1, column=2, sticky=W)

            button1 = ttk.Button(self.frame, text="Enter", command=lambda: self.multiple_data_check())
            button1.grid(row=1, column=7, sticky=W)

            button2 = ttk.Button(self.frame, text="Email", command=lambda: self.Email_open())
            button2.grid(row=1, column=8, sticky=W)

            button4 = ttk.Button(self.frame, text="Create file for all pets!", command=lambda: self.Total_file())
            button4.grid(row=1, column=9, sticky=W)

            Lable = Label(self.frame, text='Animal ID')
            Lable.grid(row=4, column=1, sticky=W)

            Lable = Label(self.frame, text='Animal Name')
            Lable.grid(row=4, column=2, sticky=W)

            Lable = Label(self.frame, text='Animal Type')
            Lable.grid(row=4, column=3, sticky=W)

            Lable = Label(self.frame, text='Animal Gender')
            Lable.grid(row=4, column=4, sticky=W)

            Lable = Label(self.frame, text='Animal Breed')
            Lable.grid(row=4, column=5, sticky=W)

            Lable = Label(self.frame, text='Animal colour')
            Lable.grid(row=4, column=6, sticky=W)

            Lable = Label(self.frame, text='Address')
            Lable.grid(row=4, column=7, sticky=W)

        except:
            tk.messagebox.showinfo("ERROR", "An error has occurred ")
    def Email_open(self):
        if admin=='True':
            window_2 = tk.Tk()
            Email(window_2).grid()
            window_2.title("Email")
            window_2.mainloop()
        else:
            tk.messagebox.showinfo("ERROR", "You do not have permission to do this")
    def keydown(self,A):
        self.user_input = Label(self.frame, textvariable = self.text_show,width=20,background='White')
        self.user_input.grid(row=1, column=2, sticky=W)
        if A.char == "":
            self.text = self.text[:len(self.text)-1]
        else:
            self.text += A.char
        print(self.text)
        self.text_show.set(self.text)
        search = self.text + '%'
        with sqlite3.connect('Data.db') as db:
            data = db.cursor()
            data.execute(
                'select * from "Animal Data" where Animal_ID LIKE ? or Animal_Name LIKE ? or Animal_Type LIKE ? or Animal_Gender LIKE ? or Animal_breed LIKE ? or Animal_colour LIKE ? or Address LIKE ?'
                ,(search,search,search,search,search,search,search,))
            self.data = data.fetchall()
            data.execute(
                'select Animal_ID from "Animal Data" where Animal_ID LIKE ?',(search,))
            self.ID=data.fetchall()
            data.execute(
                'select Animal_Name from "Animal Data" where Animal_Name LIKE ?',(search,))
            self.Animal_Name=data.fetchall()
            data.execute(
                'select Animal_Type from "Animal Data" where Animal_Type LIKE ?',(search,))
            self.Animal_Type=data.fetchall()
            data.execute(
                'select Animal_Gender from "Animal Data" where Animal_Gender LIKE ?',(search,))
            self.Animal_Gender=data.fetchall()
            data.execute(
                'select Animal_breed from "Animal Data" where Animal_breed LIKE ?',(search,))
            self.Animal_breed=data.fetchall()
            data.execute(
                'select Animal_colour from "Animal Data" where Animal_colour LIKE ?',(search,))
            self.Animal_colour=data.fetchall()
            data.execute(
                'select Address from "Animal Data" where Address LIKE ?',(search,))
            self.Address=data.fetchall()
        try:
            if len(self.ID)!=0:
                button1 = ttk.Button(self.frame, text=self.ID[0][0],width=30, command=lambda: self.update_text(self.ID[0][0]))
                button1.grid(row=2, column=1, sticky=W)
            else:
                Label_1 = Label(self.frame, text=' ',width=30)
                Label_1.grid(row=2, column=1, sticky=W)
            if len(self.Animal_Name)!=0:
                button2 = ttk.Button(self.frame, text=self.Animal_Name[0][0],width=30,command=lambda: self.update_text(self.Animal_Name[0][0]))
                button2.grid(row=2, column=2, sticky=W)
            else:
                Label_2 = Label(self.frame, text=' ',width=30)
                Label_2.grid(row=2, column=2, sticky=W)
            if len(self.Animal_Type)!=0:
                button3 = ttk.Button(self.frame, text=self.Animal_Type[0][0],width=30, command=lambda: self.update_text(self.Animal_Type[0][0]))
                button3.grid(row=2, column=3, sticky=W)
            else:
                Label_3 = Label(self.frame, text=' ',width=30)
                Label_3.grid(row=2, column=3, sticky=W)
            if len(self.Animal_Gender)!=0:
                button4 = ttk.Button(self.frame, text=self.Animal_Gender[0][0],width=30, command=lambda: self.update_text(self.Animal_Gender[0][0]))
                button4.grid(row=2, column=4, sticky=W)
            else:
                Label_4 = Label(self.frame, text=' ',width=30)
                Label_4.grid(row=2, column=4, sticky=W)
            if len(self.Animal_breed)!=0:
                button5 = ttk.Button(self.frame, text=self.Animal_breed[0][0],width=30, command=lambda: self.update_text(self.Animal_breed[0][0]))
                button5.grid(row=2, column=5, sticky=W)
            else:
                Label_5 = Label(self.frame, text=' ',width=30)
                Label_5.grid(row=2, column=5, sticky=W)
            if len(self.Animal_colour)!=0:
                button6 = ttk.Button(self.frame, text=self.Animal_colour[0][0],width=30, command=lambda: self.update_text(self.Animal_colour[0][0]))
                button6.grid(row=2, column=6, sticky=W)
            else:
                Label_6 = Label(self.frame, text=' ',width=30)
                Label_6.grid(row=2, column=6, sticky=W)
            if len(self.Address)!=0:
                button7 = ttk.Button(self.frame, text=+self.Address[0][0],width=35, command=lambda: self.update_text(self.Address[0][0]))
                button7.grid(row=2, column=7, sticky=W)
            else:
                Label_7 = Label(self.frame, text=' ',width=35)
                Label_7.grid(row=2, column=7, sticky=W)
        except IndexError:
            tk.messagebox.showinfo("ERROR", "No valid search")
            window.destroy()



    def update_text(self,text):
        self.text=text
        self.text_show.set(self.text)
        with sqlite3.connect('Data.db') as db:
            data = db.cursor()
            data.execute(
                'select * from "Animal Data" where Animal_ID = ? or Animal_Name = ? or Animal_Type = ? or Animal_Gender = ? or Animal_breed = ? or Animal_colour = ? or Address = ?'
                ,(text,text,text,text,text,text,text,))
            self.data = data.fetchall()
        self.show_data()
    def Add(self):
        self.count+=5
        if self.count>len(self.data):
            self.count=len(self.data)
        else:
            self.show_data()
        self.show_data()
        print(self.count)
    def subtract(self):
        self.count-=5
        self.show_data()
        print(self.count)
    def multiple_data_check(self):
        search_1=''
        search_2=''
        search_3=''
        search_4=''
        sec=False
        if '+' in self.text:
            text_array = list(self.text)
            for i in range(len(text_array)):
                print(str(text_array[i]))
                if text_array[i] != '+' and sec != True:
                    search_1 = search_1 + str(text_array[i])
                elif sec == True:
                    search_2 = search_2 + str(text_array[i])
                elif text_array[i] == '+':
                    sec = True
        sec = False
        if '+' in search_2:
            text_array = list(search_2)
            for i in range(len(text_array)):
                print(str(text_array[i]))
                if text_array[i] != '+' and sec != True:
                    search_4 = search_4 + str(text_array[i])
                elif sec == True:
                    search_3 = search_3 + str(text_array[i])
                elif text_array[i] == '+':
                    sec = True
        print(search_1,'     ',search_2,'    ',search_3,'    ',search_4)
        print(len(search_2))
        if len(search_2)>1:
            self.data = []
            if len(search_4) > 1:
                print('trigger')
                data_temp=[]
                with sqlite3.connect('Data.db') as db:
                    data = db.cursor()
                    data.execute(
                        'select * from "Animal Data" where Animal_ID = ? or Animal_Name = ? or Animal_Type = ? or Animal_Gender = ? or Animal_breed = ? or Animal_colour = ? or Address = ?'
                        , (search_1, search_1, search_1, search_1, search_1, search_1, search_1,))
                    self.data_1 = data.fetchall()
                    data.execute(
                        'select * from "Animal Data" where Animal_ID = ? or Animal_Name = ? or Animal_Type = ? or Animal_Gender = ? or Animal_breed = ? or Animal_colour = ? or Address = ?'
                        , (search_4, search_4, search_4, search_4, search_4, search_4, search_4,))
                    self.data_2 = data.fetchall()
                    data.execute(
                        'select * from "Animal Data" where Animal_ID = ? or Animal_Name = ? or Animal_Type = ? or Animal_Gender = ? or Animal_breed = ? or Animal_colour = ? or Address = ?'
                        , (search_3, search_3, search_3, search_3, search_3, search_3, search_3,))
                    self.data_3 = data.fetchall()
                for A in self.data_1:
                    if A in self.data_2:
                        data_temp.append(A)
                for A in self.data_3:
                    if A in data_temp and self.data_2:
                        self.data.append(A)
                print(self.data)
                self.show_data()
            else:

                with sqlite3.connect('Data.db') as db:
                    data = db.cursor()
                    data.execute(
                        'select * from "Animal Data" where Animal_ID = ? or Animal_Name = ? or Animal_Type = ? or Animal_Gender = ? or Animal_breed = ? or Animal_colour = ? or Address = ?'
                        , (search_1, search_1, search_1, search_1, search_1, search_1, search_1,))
                    self.data_1 = data.fetchall()
                    data.execute(
                        'select * from "Animal Data" where Animal_ID = ? or Animal_Name = ? or Animal_Type = ? or Animal_Gender = ? or Animal_breed = ? or Animal_colour = ? or Address = ?'
                        , (search_2, search_2, search_2, search_2, search_2, search_2, search_2,))
                    self.data_2 = data.fetchall()
                for A in self.data_1:
                    if A in self.data_2:
                        self.data.append(A)
                self.show_data()
        else:
            with sqlite3.connect('Data.db') as db:
                data = db.cursor()
                data.execute(
                    'select * from "Animal Data" where Animal_ID Like ? or Animal_Name Like ? or Animal_Type Like ? or Animal_Gender Like ? or Animal_breed Like ? or Animal_colour Like ? or Address LIKE ?'
                    , (search_1, search_1, search_1, search_1, search_1, search_1, search_1,))
                self.data = data.fetchall()
            self.show_data()

    def show_data(self):

        button1 = ttk.Button(self.frame, text='Next 5', command=lambda: self.Add())
        button1.grid(row=3, column=2, sticky=W)
        button2 = ttk.Button(self.frame, text='previous 5', command=lambda: self.subtract())
        button2.grid(row=3, column=1, sticky=W)
        if len(self.data)>0:
            y=5
            for i in range(self.count-5,self.count):
                if i==len(self.data):
                    break
                else:
                    try:
                        photo = ImageTk.PhotoImage(file="Images/"+self.data[i][7]+'.jpg')
                        result_img = Label(self.frame, image=photo)
                        result_img.image = photo
                        result_img.grid(row=y, column=8, sticky=W)
                    except:
                        tk.messagebox.showinfo("ERROR", "An error has occurred with loading image")
                    result0 = Label(self.frame, text=self.data[i][0],width=15,background='White')
                    result0.grid(row=y, column=1, sticky=W)
                    result1 = Label(self.frame, text=self.data[i][1],width=20,background='White')
                    result1.grid(row=y, column=2, sticky=W)
                    result2 = Label(self.frame, text=self.data[i][2],width=15,background='White')
                    result2.grid(row=y, column=3, sticky=W)
                    result3 = Label(self.frame, text=self.data[i][3],width=15,background='White')
                    result3.grid(row=y, column=4, sticky=W)
                    result4 = Label(self.frame, text=self.data[i][4],width=25,background='White')
                    result4.grid(row=y, column=5, sticky=W)
                    result5 = Label(self.frame, text=self.data[i][5],width=25,background='White')
                    result5.grid(row=y, column=6, sticky=W)
                    result6 = Label(self.frame, text=self.data[i][6],width=35,background='White')
                    result6.grid(row=y, column=7, sticky=W)
                    y+=1
        else:
            tk.messagebox.showinfo("Missing Data", "Please enter a valid search")
            window.destroy()
        self.text=''
        self.text_show=StringVar()
        self.text_show.set(self.text)
        self.bind("<KeyPress>",self.keydown)
        self.grid()
        self.focus_set()
    def Total_file(self):
        try:
            document = Document()
            document.add_picture('LOGO.png')
            document.add_heading('Here are some awesome Pets to find')
            p = document.add_paragraph('Below is a table containing All animals currently up for adoption!')
            with sqlite3.connect('Data.db') as db:
                data = db.cursor()
                data.execute(
                    'select Animal_ID, Animal_Name, Animal_Type, Animal_Gender, Animal_breed, Animal_colour, Address from "Animal Data"')
            data = data.fetchall()
            table = document.add_table(rows=1, cols=11)
            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Animal ID'
            hdr_cells[1].text = 'Animal Name'
            hdr_cells[2].text = 'Animal Type'
            hdr_cells[3].text = 'Animal Gender'
            hdr_cells[4].text = 'Animal breed'
            hdr_cells[5].text = 'Animal colour'
            hdr_cells[6].text = 'address'
            for Animal_ID, Animal_Name, Animal_Type, Animal_Gender, Animal_breed, Animal_colour, Address in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(Animal_ID)
                row_cells[1].text = str(Animal_Name)
                row_cells[2].text = str(Animal_Type)
                row_cells[3].text = str(Animal_Gender)
                row_cells[4].text = str(Animal_breed)
                row_cells[5].text = str(Animal_colour)
                row_cells[6].text = str(Address)
            document.add_page_break()
            document.save('Animal Place Pets.docx')
            tk.messagebox.showinfo("success", "success you now have all of our pets information ")
        except:
            tk.messagebox.showinfo("ERROR", "There has unfortunately been an Error")
class create_user(tk.Frame):
    def __init__(self, parent):
        tk.Frame.__init__(self, parent)
        Lable= tk.Label(self, text='Join Animal Place and get updates when new pets are ready for adoption')
        Lable.grid(row=1, column=1, sticky=W)

        Lable= tk.Label(self, text='First Name:')
        Lable.grid(row=2, column=1, sticky=W)

        self.First = ttk.Entry(self)
        self.First.grid(row=2, column=2, sticky=W)

        Lable= tk.Label(self, text='Last Name:')
        Lable.grid(row=3, column=1, sticky=W)

        self.Last = ttk.Entry(self)
        self.Last.grid(row=3, column=2, sticky=W)

        Lable= tk.Label(self, text='Email:')
        Lable.grid(row=4, column=1, sticky=W)

        self.Email = ttk.Entry(self)
        self.Email.grid(row=4, column=2, sticky=W)

        Lable = tk.Label(self, text='Password:')
        Lable.grid(row=5, column=1, sticky=W)

        self.Password = ttk.Entry(self)
        self.Password.grid(row=5, column=2, sticky=W)

        button1 = ttk.Button(self, text="JOIN!!", command=lambda: self.Join())
        button1.grid(row=6, column=2, sticky=W)
    def Join(self):
        self.First=self.First.get()
        self.Last=self.Last.get()
        self.Email=self.Email.get()
        self.Password = self.Password.get()
        if self.First=='' or self.Last=='' or self.Email=='' or self.Password=='':
            tk.messagebox.showinfo("ERROR", "We are missing some data please enter all required data")
            self.destroy()
            Login.create_user_open(self)
        else:
            self.Password=hashlib.md5(self.Password.encode('utf-8')).hexdigest()
            with sqlite3.connect('Data.db') as db:
                data = db.cursor()
                data.execute('select User_ID from "User Data"')
            self.ID=len(data.fetchall())+1
            with sqlite3.connect('Data.db') as db:
                data = db.cursor()
                data.execute('INSERT INTO "User Data" values (?,?,?,?,?,"False")',(self.ID,self.Password,self.First,self.Last,self.Email))
            tk.messagebox.showinfo("Success", "WELCOME TO THE COMMUNITY YOUR ID IS "+str(self.ID))
            self.destroy()

class Email(tk.Frame):
    def __init__(self, parent):
        tk.Frame.__init__(self, parent)
                                       
        Lable= tk.Label(self, text='Let others Know about your pet')
        Lable.grid(row=1, column=1, sticky=W)

        Lable= tk.Label(self, text='Animal Name')
        Lable.grid(row=2, column=1, sticky=W)
                                       
        self.Name = ttk.Entry(self)
        self.Name.grid(row=2, column=2, sticky=W)

        Lable= tk.Label(self, text='Animal Type')
        Lable.grid(row=3, column=1, sticky=W)
                                       
        self.Type = ttk.Entry(self)
        self.Type.grid(row=3, column=2, sticky=W)

        Lable= tk.Label(self, text='Animal Gender')
        Lable.grid(row=4, column=1, sticky=W)
                                       
        self.Gender = ttk.Entry(self)
        self.Gender.grid(row=4, column=2, sticky=W)

        Lable= tk.Label(self, text='Animal breed')
        Lable.grid(row=5, column=1, sticky=W)
                                       
        self.Breed = ttk.Entry(self)
        self.Breed.grid(row=5, column=2, sticky=W)

        Lable= tk.Label(self, text='Animal colour')
        Lable.grid(row=6, column=1, sticky=W)
                                       
        self.colour = ttk.Entry(self)
        self.colour.grid(row=6, column=2, sticky=W)

        Lable= tk.Label(self, text='Address')
        Lable.grid(row=7, column=1, sticky=W)
                                       
        self.Address = ttk.Entry(self)
        self.Address.grid(row=7, column=2, sticky=W)

        button1 = ttk.Button(self, text="Send", command=lambda: self.send_mail())
        button1.grid(row=8, column=2, sticky=W)
        with sqlite3.connect('Data.db') as db:
                Data = db.cursor()
                Data.execute(
                    'select Email from "User Data"')
                self.Emails = Data.fetchall()

    def send_mail(self):
        email = 'duglassmick@gmail.com'
        password = 'Python123'
        content = 'New animal Everyone meet '+self.Name.get()+" they're a "+self.Type.get()+' there geneder is '+self.Gender.get()+' colour '+self.colour.get()+' if you would like this animal there located at '+self.Address.get()
        
        for i in range(0,len(self.Emails)):
            to_email=self.Emails[0][i]
            try:
                mail = smtplib.SMTP('smtp.gmail.com', 587)
                mail.ehlo()
                mail.starttls()
                mail.login(email,password)
                mail.sendmail(email, to_email,content)
                mail.close()
                tk.messagebox.showinfo("Success", "Message was sent")
            except:
                tk.messagebox.showinfo("ERROR", "An error has occurred")

open = Tk()
open.title('Animal Place')
Login(open).grid()
open.mainloop()
